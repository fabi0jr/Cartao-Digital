import re
import time
import PyPDF2
import  os
import datetime
import pandas as pd
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
import shutil  # Importe o módulo shutil para mover arquivos
from selenium.webdriver.support.ui import Select
from docx import Document
from docx.shared import Pt
import pyautogui
from PIL import Image, ImageDraw, ImageFont
import textwrap
import pytz  # Importe a biblioteca pytz para lidar com fusos horários
import openpyxl
import json

def carregar_senha():
    try:
        with open(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\senha.json", "r") as file:
            data = json.load(file)
            return data.get("senha")
    except (FileNotFoundError, json.JSONDecodeError):
        return None

def salvar_nova_senha(nova_senha):
    data = {"senha": nova_senha}
    with open(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\senha.json", "w") as file:
        json.dump(data, file)

def formatar_cpf(cpf):
    # Remove qualquer caractere não numérico do CPF
    cpf = re.sub(r'\D', '', cpf)
    
    # Verifica se o CPF tem 11 dígitos
    if len(cpf) != 11:
        print("CPF inválido. Certifique-se de digitar 11 dígitos.")
        return None
    
    # Formata o CPF como "xxx.xxx.xxx-xx"
    cpf_formatado = f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"
    return cpf_formatado

def inserir_cpfs():
        
    senha = carregar_senha() 
  # Configurar o navegador em modo headless
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--disable-notifications')
    options.add_argument('--window-size=1920,1080')
    driver = webdriver.Chrome(options=options)

    # Realizar login no site
    driver.get('http://sistemas.sefaz.am.gov.br/gcc/entrada.do')  # Substitua pelo URL do site
    time.sleep(3)
    usuario = driver.find_element(By.ID, 'username')  # Substitua pelo campo de usuário
    senha_da_pagina = driver.find_element(By.ID, 'password')  # Substitua pelo campo de senha
    botao_login = driver.find_element(By.XPATH, '//*[@id="fm1"]/fieldset/div[3]/div/div[4]/input[4]')  # Substitua pelo botão de login

    usuario.send_keys('03483401253')
    senha_da_pagina.send_keys(senha)
    botao_login.click()

    while True:
        cpf = input("Insira um CPF (ou apenas pressione Enter para parar): ")
        if not cpf:
            break

        # Verificar se o CPF já existe na planilha
        cpf_formatado = formatar_cpf(cpf)
        planilha = pd.read_excel(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx")

        if cpf_formatado in planilha['CPF'].values:
            print(f"CPF {cpf_formatado} já existe na planilha. Ignorando...")
            continue

        elemento = driver.find_element(By.XPATH,'//*[@id="oCMenu___GCC2300"]')
        elemento.click()

        dropdown_div = driver.find_element(By.XPATH, '//*[@id="oCMenu___GCC1008"]')
        dropdown_div.click()

        caixacpf = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')
        caixacpf.click()
        caixacpf.send_keys(cpf)

        consultar = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_cadastroProdutorRuralAction!pesquisarProdutorRural"]')
        consultar.click()

        try:
            comcpf = driver.find_element(By.XPATH, '//*[@id="tbProdutorRural"]/tbody/tr/td[8]/a[2]/img')
            planilha = pd.read_excel(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx")
            novo_dataframe = pd.DataFrame({'CPF': [cpf_formatado]})
            planilha = pd.concat([planilha, novo_dataframe], ignore_index=True)
            planilha.to_excel(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx", index=False)
            print(f"CPF {cpf_formatado} inserido com sucesso!")

        except NoSuchElementException:
            print(f"CPF {cpf_formatado} não encontrado")

    driver.quit()

def emitir_cartoes():
    senha = carregar_senha()
        # Configurar o navegador em modo headless
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--disable-notifications')
    options.add_argument('--window-size=1920,1080')
    driver = webdriver.Chrome(options=options)

    # Realizar login no site
    driver.get('http://sistemas.sefaz.am.gov.br/gcc/entrada.do')  # Substitua pelo URL do site
    time.sleep(3)
    usuario = driver.find_element(By.ID, 'username')  # Substitua pelo campo de usuário
    senha_da_pagina = driver.find_element(By.ID, 'password')  # Substitua pelo campo de senha
    botao_login = driver.find_element(By.XPATH, '//*[@id="fm1"]/fieldset/div[3]/div/div[4]/input[4]')  # Substitua pelo botão de login

    usuario.send_keys('03483401253')
    senha_da_pagina.send_keys(senha)
    botao_login.click()
  
    while True:
        planilha = pd.read_excel(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx")
        # Crie um DataFrame vazio para armazenar os dados
        # Crie um arquivo Excel para armazenar os dados
        os.path.isfile(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")
        workbook = openpyxl.load_workbook(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")
        sheet = workbook.active
        for index, row in planilha.iterrows():
            cpf = row['CPF']  # Pegar o CPF da linha atual

            # Configurar o fuso horário de Manaus (GMT-4)
            fuso_horario = pytz.timezone('America/Manaus')
            # Obtém a data atual no fuso horário de Manaus
            data_atual_manaus = datetime.datetime.now(fuso_horario)
            largura_maxima_x = 540  # Ajuste conforme necessário
            # Abrir o modelo PNG
            modelo = Image.open(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\img\frente.png")
            modelo_verso = Image.open(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\img\verso.png")
            fonte = ImageFont.truetype(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 41)
            fonte_endereco = ImageFont.truetype(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 38)
            fonte_atv2 = ImageFont.truetype(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 39)
            desenho = ImageDraw.Draw(modelo)
            desenho_verso = ImageDraw.Draw(modelo_verso)

            while True:
                elemento = driver.find_element(By.XPATH,'//*[@id="oCMenu___GCC2300"]')
                elemento.click()
                
                dropdown_div = driver.find_element(By.XPATH, '//*[@id="oCMenu___GCC1008"]')
                dropdown_div.click()
                caixacpf = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')  
                caixacpf.click()  
                caixacpf.send_keys(cpf)  
                consultar = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_cadastroProdutorRuralAction!pesquisarProdutorRural"]')       
                consultar.click()
                time.sleep(1)
                situacao = driver.find_element(By.XPATH, '//*[@id="tbProdutorRural"]/thead/tr/th[2]')
                situacao.click()
                time.sleep(2)
                try:
                    abadeclaracao = WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.XPATH, '//*[@id="tbProdutorRural"]/tbody/tr[1]/td[8]/a[2]'))
                    )
                    abadeclaracao.click()
                    # Resto do seu código
                except Exception as e:
                    print(f"Erro ao clicar em 'abadeclaracao' para o CPF {cpf}: {str(e)}")
                  
                nome_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cceaPessoaFisica_pfNome"]')
                # Use JavaScript para obter o valor do atributo 'value' do elemento
                nome = driver.execute_script("return arguments[0].value;", nome_element)

                rp_da_pagina = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_ieProdutorRuralFormatado"]')
                rp = driver.execute_script("return arguments[0]. value;", rp_da_pagina)

                cpf_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')
                cpf = driver.execute_script("return arguments[0]. value;", cpf_element)

                propiedade_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmPropriedade"]')
                propriedade = driver.execute_script("return arguments[0]. value;", propiedade_element)

                endereco_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_txEnderecoPropriedade"]')
                endereco = driver.execute_script("return arguments[0]. value;", endereco_element)

                unloc_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_sgDistritoIdam"]')
                unloc_da_pagina = driver.execute_script("return arguments[0]. value;", unloc_element)
                
                if unloc_da_pagina == 'BAE':
                    unloc_da_pagina = 'BAR'
                
                mapeamento_municipios = {
                    'ALV': 'Alvaraes',
                    'AMT': 'Amatura',
                    'ANA':'Anamã',
                    'ANO':'Anori',          
                    'APU':'Apuí',
                    'ATN':'Atalaia do Norte',
                    'ATZ':'Autazes',
                    'BAZ':'Barcelos',
                    'BAE':'Barreirinha',
                    'BAR':'Barreirinha',
                    'BJC':'Benjamin Constant',
                    'BER':'Beruri',
                    'BVR':'Boa Vista do Ramos',
                    'BOA':'Boca do Acre',
                    'BBA':'Borba',
                    'CAP':'Caapiranga',
                    'CAN':'Canutama',
                    'CAF':'Carauari',
                    'CAR':'Careiro',
                    'CAZ':'Careiro da Várzea',
                    'CRZ':'Coari',
                    'COD':'Codajás',
                    'ERN':'Eirunepé',
                    'EUR':'Envira',
                    'FBA':'Fonte Boa',
                    'GAJ':'Guajará',
                    'HIA':'Humaitá',
                    'IPX':'Ipixuna',
                    'IRB':'Iranduba',
                    'ITA':'Itamarati',
                    'ITR':'Itacoatiara',
                    'ITG':'Itapiranga',
                    'JPR':'Japurá',
                    'JUR':'Juruá',
                    'JUT':'Jutaí',
                    'LBR':'Lábrea',
                    'MPU':'Manacapuru',
                    'MQR':'Manaquiri',
                    'MAO':'Manaus',
                    'MNX':'Manicoré',
                    'MTS':'Monte Sinai',
                    'MRA':'Maraã',
                    'MBZ':'Maués',
                    'NMD':'Nhamundá',
                    'NRO':'Novo Remanso',
                    'MTP':'Santo Antonio do Matupi',
                    'VE':'Vila Extrema',
                    'VRC':'Vila Rica de Caviana',
                    'BNA':'Balbina',
                    'VLD':'Vila de Lindoia',
                    'RLD':'Vila da Realidade',
                    'NON':'Nova Olinda do Norte',
                    'NAR':'Novo Airão',
                    'NAP':'Novo Aripuanã',
                    'PAR':'Parintins',
                    'PUI':'Pauini',
                    'PRF':'Presidente Figueiredo',
                    'RPE':'Rio Preto da Eva',
                    'SIR':'Santa Isabel do Rio Negro',
                    'SAI':'Santo Antônio do Içá',
                    'SJL':'São Gabriel da Cachoeira',
                    'SPO':'São Paulo de Olivença',
                    'SSU':'São Sebastião do Uatumã',
                    'SUL-CAN':'Sul de Canutama',
                    'SLV':'Silves',
                    'TBT':'Tabatinga',
                    'TPA':'Tapauá',
                    'TFF':'Tefé',
                    'TNT':'Tonantins',
                    'UAN':'Uarini',
                    'URC':'Urucará',
                    'UCB':'Urucurituba',
                }

                valor_numerico = unloc_da_pagina

                nome_municipio = mapeamento_municipios.get(valor_numerico, 'MUNICIPIO_DESCONHECIDO')

                latitude_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLatitude"]')
                latitude = driver.execute_script("return arguments[0]. value;", latitude_element)

                longitude_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLongitude"]')
                longitude = driver.execute_script("return arguments[0]. value;", longitude_element)

                atv1_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaePrincipal"]')
                atv1= driver.execute_script("return arguments[0]. value;", atv1_element)

                atv2_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaeSecundario"]')
                atv2 = driver.execute_script("return arguments[0]. value;", atv2_element)

                inicioatv_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_anoInicioAtividade"]')
                inicioatv = driver.execute_script("return arguments[0]. value;", inicioatv_element)

                numcontrole_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nrDeclaracaoUnidLocal"]')
                numcontrole = driver.execute_script("return arguments[0]. value;",numcontrole_element)

                cnae1_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cnaePrincipalFormatado"]')
                cnae1 = driver.execute_script("return arguments[0]. value;",cnae1_element)

                cnae2_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cnaeSecundarioFormatado"]')
                cnae2 = driver.execute_script("return arguments[0]. value;",cnae2_element)
                
                if not atv2:
                    cnae2 = ""
                  
                data_atual_manaus_str = data_atual_manaus.strftime("%d/%m/%Y")
                nova_linha = [nome, cpf, nome_municipio, data_atual_manaus_str, cnae1, cnae2]
                sheet.append(nova_linha)
                unloc = "PR-" + unloc_da_pagina + "/" + numcontrole
                unloc = str(unloc)
                validade = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_dtValidadeDeclaracaoFormatado"]')
                validade_da_pagina = driver.execute_script("return arguments[0]. value;",validade)
                print(nome, "|", unloc, "|")
                pasta_unloc = os.path.join(r"I:\CARTEIRAS DIGITAIS", unloc_da_pagina)
                
                if not os.path.exists(pasta_unloc):
                    # Se a pasta não existe, crie-a
                    os.mkdir(pasta_unloc)
                  
                data_atual = datetime.datetime.now()
                pasta_data_atual = os.path.join(pasta_unloc, data_atual.strftime("%d.%m.%Y"))

                if not os.path.exists(pasta_data_atual):
                    # Se a pasta da data atual não existe, crie-a
                    os.mkdir(pasta_data_atual)
                        
                # Caminho completo para salvar o PDF na pasta da data atual
                pdf_path = os.path.join(pasta_data_atual, nome + '.pdf')

                # Defina a pasta da data atual como a pasta para salvar o PDF
                output_pdf_filename = pdf_path

                coordenadas_frente = {
                    "RP": (217,393),
                    "NOME": (95,518,1309,604),
                    "CPF": (864,392),
                    "NOME PROPRIEDADE": (100, 660),
                    "UNLOC": (212,824),
                    "INICIO ATV": (751,824),
                    "VALIDADE": (1063,825),
                }

                # Inserir informações no modelo nas coordenadas especificadas
                desenho.text(coordenadas_frente["RP"], rp, fill=(0,0,0), font=fonte)
                desenho.text(coordenadas_frente["NOME"], nome, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["CPF"], cpf, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["NOME PROPRIEDADE"], propriedade, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["UNLOC"], unloc, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["INICIO ATV"], inicioatv, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["VALIDADE"], validade_da_pagina, fill=(0, 0, 0), font=fonte)

                # Quebra o texto em várias linhas
                linhas_endereco = textwrap.wrap(endereco, largura_maxima_x)
                def limitar_texto(texto, comprimento_maximo):
                    if len(texto) > comprimento_maximo:
                            return texto[:comprimento_maximo - 3]
                    else:
                        return texto
                endereco = limitar_texto(endereco, 50)

                linhas_atv2 = textwrap.wrap(atv2, largura_maxima_x)
                def limitar_texto(texto, comprimento_maximo):
                    if len(texto) > comprimento_maximo:
                            return texto[:comprimento_maximo - 3]
                    else:
                        return texto
                atv2 = limitar_texto(atv2, 60)
                
                coordenadas_verso = {
                    "END": (89,285),
                    "ATV1": (89,473),
                    "ATV2": (89,631),
                    "LOC": (382,804),
                }

                def desenhar_texto_quebrado(coordenadas, texto, desenho, fonte, largura_maxima_x):
                    coordenada = (coordenadas[0], coordenadas[1])
                    linhas = textwrap.wrap(texto, width=largura_maxima_x // 9)
                    for linha in linhas:
                        linhas_quebradas = textwrap.wrap(linha, width=largura_maxima_x // 9)

                        for linha_quebrada in linhas_quebradas:
                            if coordenada[1] < modelo_verso.size[1]:
                                desenho.text(coordenada, linha_quebrada, fill=(0, 0, 0), font=fonte)
                                coordenada = (coordenada[0], coordenada[1] + 40)  # Aumente o valor (ex: +50) conforme necessário
                            else:
                                break

                desenho_verso.text(coordenadas_verso["END"], endereco, fill=(0, 0, 0), font=fonte_endereco)

                coordenada_endereco = (coordenadas_verso["END"][0], coordenadas_verso["END"][1])

                # Loop para desenhar cada linha do endereço
                for linha in linhas_endereco:
                    linhas_quebradas = textwrap.wrap(linha, width=largura_maxima_x // 9)

                    for linha_quebrada in linhas_quebradas:
                        if coordenada_endereco[1] < modelo_verso.size[1]:
                            desenho_verso.text(coordenada_endereco, linha_quebrada, fill=(0, 0, 0), font=fonte_endereco)
                            coordenada_endereco = (coordenada_endereco[0], coordenada_endereco[1] + 40)  # Aumente o valor (ex: +50) conforme necessário
                        else:
                            break

                coordenada_atv2 = (coordenadas_verso["ATV2"][0], coordenadas_verso["ATV2"][1])

                desenhar_texto_quebrado(coordenadas_verso["END"], endereco, desenho_verso, fonte_endereco, largura_maxima_x)
                desenhar_texto_quebrado(coordenadas_verso["ATV1"], cnae1 + " - " + atv1, desenho_verso, fonte, largura_maxima_x)
                if atv2:
                    desenhar_texto_quebrado(coordenadas_verso["ATV2"], cnae2 + " - " + atv2, desenho_verso, fonte, largura_maxima_x)
                desenhar_texto_quebrado(coordenadas_verso["LOC"], latitude + "  " + longitude, desenho_verso, fonte, largura_maxima_x)

                modelo.save(cpf + ".pdf")
                modelo_verso.save("verso.pdf")
                # Nome dos arquivos PDF que você gerou
                pdf1_filename = cpf + ".pdf"
                pdf2_filename = "verso.pdf"

                # Abra os arquivos PDF
                pdf1 = PyPDF2.PdfReader(pdf1_filename)
                pdf2 = PyPDF2.PdfReader(pdf2_filename)

                # Crie um objeto PDFFileMerger para mesclar os PDFs
                pdf_merger = PyPDF2.PdfMerger()

                # Adicione os arquivos PDF à mesclagem
                pdf_merger.append(pdf1)
                pdf_merger.append(pdf2)

                output_pdf_filename = os.path.join(r"I:\CARTEIRAS DIGITAIS",unloc_da_pagina, pasta_data_atual , nome  + '.pdf')
                pdf_merger.write(output_pdf_filename)

                pdf_merger.close()
                os.remove(cpf + ".pdf")
                os.remove("verso.pdf")
                planilha = planilha[planilha['CPF'] != cpf]

                break
        workbook.save('banco_de_dados.xlsx')
        # Salvar a planilha atualizada após a exclusão do CPF
        planilha.to_excel('cpfs.xlsx', index=False)

        break

def emitir_cartoes_slash():

    senha = carregar_senha()
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--disable-notifications')
    options.add_argument('--window-size=1920,1080')
    driver = webdriver.Chrome(options=options)

    # Realizar login no site
    driver.get('http://sistemas.sefaz.am.gov.br/gcc/entrada.do')  # Substitua pelo URL do site
    time.sleep(3)
    usuario = driver.find_element(By.ID, 'username')  # Substitua pelo campo de usuário
    senha_da_pagina = driver.find_element(By.ID, 'password')  # Substitua pelo campo de senha
    botao_login = driver.find_element(By.XPATH, '//*[@id="fm1"]/fieldset/div[3]/div/div[4]/input[4]')  # Substitua pelo botão de login

    usuario.send_keys('03483401253')
    senha_da_pagina.send_keys(senha)
    botao_login.click()


    while True:
        planilha = pd.read_excel(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx")
        # Crie um DataFrame vazio para armazenar os dados
        # Crie um arquivo Excel para armazenar os dados
        os.path.isfile(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")
        workbook = openpyxl.load_workbook(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")
        sheet = workbook.active

        for index, row in planilha.iterrows():
            cpf = row['CPF']  # Pegar o CPF da linha atual

            # Configurar o fuso horário de Manaus (GMT-4)
            fuso_horario = pytz.timezone('America/Manaus')
            # Obtém a data atual no fuso horário de Manaus
            data_atual_manaus = datetime.datetime.now(fuso_horario)

            largura_maxima_x = 540  # Ajuste conforme necessário
          
            modelo = Image.open(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\img\frente.png")
            modelo_verso = Image.open(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\img\verso.png")

            fonte = ImageFont.truetype(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 41)
            fonte_endereco = ImageFont.truetype(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 38)
            fonte_atv2 = ImageFont.truetype(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 39)

            desenho = ImageDraw.Draw(modelo)
            desenho_verso = ImageDraw.Draw(modelo_verso)

            while True:
                
                elemento = driver.find_element(By.XPATH,'//*[@id="oCMenu___GCC2300"]')    
                elemento.click()                   
                dropdown_div = driver.find_element(By.XPATH, '//*[@id="oCMenu___GCC1008"]')
                dropdown_div.click()
                caixacpf = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')  
                caixacpf.click() 
                caixacpf.send_keys(cpf)  
                consultar = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_cadastroProdutorRuralAction!pesquisarProdutorRural"]')       
                consultar.click()
                time.sleep(1)
                situacao = driver.find_element(By.XPATH, '//*[@id="tbProdutorRural"]/thead/tr/th[2]')
                situacao.click()
                time.sleep(2)
                try:
                    abadeclaracao = WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.XPATH, '//*[@id="tbProdutorRural"]/tbody/tr[1]/td[8]/a[2]'))
                    )
                    abadeclaracao.click()
                    # Resto do seu código
                except Exception as e:
                    print(f"Erro ao clicar em 'abadeclaracao' para o CPF {cpf}: {str(e)}")
                nome_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cceaPessoaFisica_pfNome"]')
                # Use JavaScript para obter o valor do atributo 'value' do elemento
                nome = driver.execute_script("return arguments[0].value;", nome_element)

                rp_da_pagina = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_ieProdutorRuralFormatado"]')
                rp = driver.execute_script("return arguments[0]. value;", rp_da_pagina)

                cpf_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')
                cpf = driver.execute_script("return arguments[0]. value;", cpf_element)

                propiedade_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmPropriedade"]')
                propriedade = driver.execute_script("return arguments[0]. value;", propiedade_element)

                endereco_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_txEnderecoPropriedade"]')
                endereco = driver.execute_script("return arguments[0]. value;", endereco_element)

                unloc_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_sgDistritoIdam"]')
                unloc_da_pagina = driver.execute_script("return arguments[0]. value;", unloc_element)
                if unloc_da_pagina == 'BAE':
                    unloc_da_pagina = 'BAR'
                  
                mapeamento_municipios = {
                    'ALV': 'Alvaraes',
                    'AMT': 'Amatura',
                    'ANA':'Anamã',
                    'ANO':'Anori',          
                    'APU':'Apuí',
                    'ATN':'Atalaia do Norte',
                    'ATZ':'Autazes',
                    'BAZ':'Barcelos',
                    'BAE':'Barreirinha',
                    'BAR':'Barreirinha',
                    'BJC':'Benjamin Constant',
                    'BER':'Beruri',
                    'BVR':'Boa Vista do Ramos',
                    'BOA':'Boca do Acre',
                    'BBA':'Borba',
                    'CAP':'Caapiranga',
                    'CAN':'Canutama',
                    'CAF':'Carauari',
                    'CAR':'Careiro',
                    'CAZ':'Careiro da Várzea',
                    'CRZ':'Coari',
                    'CIZ':'Coari',
                    'COD':'Codajás',
                    'ERN':'Eirunepé',
                    'ENV':'Envira',
                    'FBA':'Fonte Boa',
                    'GAJ':'Guajará',
                    'HIA':'Humaitá',
                    'IPX':'Ipixuna',
                    'IRB':'Iranduba',
                    'ITA':'Itamarati',
                    'ITR':'Itacoatiara',
                    'ITG':'Itapiranga',
                    'JPR':'Japurá',
                    'JUR':'Juruá',
                    'JUT':'Jutaí',
                    'LBR':'Lábrea',
                    'MPU':'Manacapuru',
                    'MQR':'Manaquiri',
                    'MAO':'Manaus',
                    'MNX':'Manicoré',
                    'MTS-ATZ':'Monte Sinai',
                    'MRA':'Maraã',
                    'MBZ':'Maués',
                    'NMD':'Nhamundá',
                    'ITR-NRO':'Novo Remanso',
                    'MNX-MTP':'Santo Antonio do Matupi',
                    'MTP': 'Santo Antonio do Matupi',
                    'VE':'Vila Extrema',
                    'LBR-VE':'Vila Extrema',
                    'VRC':'Vila Rica de Caviana',
                    'PRF-BNA':'Balbina',
                    'VLD':'Vila de Lindoia',
                    'RLD':'Vila da Realidade',
                    'NON':'Nova Olinda do Norte',
                    'NAR':'Novo Airão',
                    'NAP':'Novo Aripuanã',
                    'PAR':'Parintins',
                    'PUI':'Pauini',
                    'PRF':'Presidente Figueiredo',
                    'RPE':'Rio Preto da Eva',
                    'SIR':'Santa Isabel do Rio Negro',
                    'SAI':'Santo Antônio do Içá',
                    'SJL':'São Gabriel da Cachoeira',
                    'SPO':'São Paulo de Olivença',
                    'SSU':'São Sebastião do Uatumã',
                    'SUL-CAN':'Sul de Canutama',
                    'SLV':'Silves',
                    'TBT':'Tabatinga',
                    'TPA':'Tapauá',
                    'TFF':'Tefé',
                    'TNT':'Tonantins',
                    'UAN':'Uarini',
                    'URC':'Urucará',
                    'UCB':'Urucurituba',
                }

                valor_numerico = unloc_da_pagina

                # Obtém o nome do município com base no valor numérico usando o mapeamento
                nome_municipio = mapeamento_municipios.get(valor_numerico, 'MUNICIPIO_DESCONHECIDO')

                latitude_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLatitude"]')
                latitude = driver.execute_script("return arguments[0]. value;", latitude_element)

                longitude_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLongitude"]')
                longitude = driver.execute_script("return arguments[0]. value;", longitude_element)

                atv1_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaePrincipal"]')
                atv1= driver.execute_script("return arguments[0]. value;", atv1_element)

                atv2_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaeSecundario"]')
                atv2 = driver.execute_script("return arguments[0]. value;", atv2_element)

                inicioatv_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_anoInicioAtividade"]')
                inicioatv = driver.execute_script("return arguments[0]. value;", inicioatv_element)

                numcontrole_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nrDeclaracaoUnidLocal"]')
                numcontrole = driver.execute_script("return arguments[0]. value;",numcontrole_element)

                cnae1_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cnaePrincipalFormatado"]')
                cnae1 = driver.execute_script("return arguments[0]. value;",cnae1_element)

                cnae2_element = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cnaeSecundarioFormatado"]')
                cnae2 = driver.execute_script("return arguments[0]. value;",cnae2_element)
                
                if not atv2:
                    cnae2 = ""
            
                descricao = input("Motivo pelo qual a carteira esta sendo emitida:")
                data_atual_manaus_str = data_atual_manaus.strftime("%d/%m/%Y")
                nova_linha = [nome, cpf, nome_municipio, data_atual_manaus_str, cnae1, cnae2, descricao]
                sheet.append(nova_linha)
                
                unloc = "PR-" + unloc_da_pagina + "/" + numcontrole
                unloc = str(unloc)
                validade = driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_dtValidadeDeclaracaoFormatado"]')
                validade_da_pagina = driver.execute_script("return arguments[0]. value;",validade)
                print(nome, "|", unloc, "|")
                pasta_unloc = os.path.join(r"I:\CARTEIRAS DIGITAIS", unloc_da_pagina)                
                if not os.path.exists(pasta_unloc):
                    # Se a pasta não existe, crie-a
                    os.mkdir(pasta_unloc)

                data_atual = datetime.datetime.now()
                pasta_data_atual = os.path.join(pasta_unloc, data_atual.strftime("%d.%m.%Y"))

                if not os.path.exists(pasta_data_atual):
                    # Se a pasta da data atual não existe, crie-a
                    os.mkdir(pasta_data_atual)
                pdf_path = os.path.join(pasta_data_atual, nome + '.pdf')
                output_pdf_filename = pdf_path
                coordenadas_frente = {
                    "RP": (217,393),
                    "NOME": (95,518,1309,604),
                    "CPF": (864,392),
                    "NOME PROPRIEDADE": (100, 660),
                    "UNLOC": (212,824),
                    "INICIO ATV": (751,824),
                    "VALIDADE": (1063,825),
                }

                # Inserir informações no modelo nas coordenadas especificadas
                desenho.text(coordenadas_frente["RP"], rp, fill=(0,0,0), font=fonte)
                desenho.text(coordenadas_frente["NOME"], nome, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["CPF"], cpf, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["NOME PROPRIEDADE"], propriedade, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["UNLOC"], unloc, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["INICIO ATV"], inicioatv, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["VALIDADE"], validade_da_pagina, fill=(0, 0, 0), font=fonte)

                # Quebra o texto em várias linhas
                linhas_endereco = textwrap.wrap(endereco, largura_maxima_x)
                def limitar_texto(texto, comprimento_maximo):
                    if len(texto) > comprimento_maximo:
                            return texto[:comprimento_maximo - 3]
                    else:
                        return texto
                endereco = limitar_texto(endereco, 50)

                linhas_atv2 = textwrap.wrap(atv2, largura_maxima_x)
                def limitar_texto(texto, comprimento_maximo):
                    if len(texto) > comprimento_maximo:
                            return texto[:comprimento_maximo - 3]
                    else:
                        return texto
                atv2 = limitar_texto(atv2, 60)
                
                coordenadas_verso = {
                    "END": (89,285),
                    "ATV1": (89,473),
                    "ATV2": (89,631),
                    "LOC": (382,804),
                }

                def desenhar_texto_quebrado(coordenadas, texto, desenho, fonte, largura_maxima_x):
                    coordenada = (coordenadas[0], coordenadas[1])
                    linhas = textwrap.wrap(texto, width=largura_maxima_x // 9)
                    for linha in linhas:
                        linhas_quebradas = textwrap.wrap(linha, width=largura_maxima_x // 9)

                        for linha_quebrada in linhas_quebradas:
                            if coordenada[1] < modelo_verso.size[1]:
                                desenho.text(coordenada, linha_quebrada, fill=(0, 0, 0), font=fonte)
                                coordenada = (coordenada[0], coordenada[1] + 40)  # Aumente o valor (ex: +50) conforme necessário
                            else:
                                break

                

                desenho_verso.text(coordenadas_verso["END"], endereco, fill=(0, 0, 0), font=fonte_endereco)

                coordenada_endereco = (coordenadas_verso["END"][0], coordenadas_verso["END"][1])

                # Loop para desenhar cada linha do endereço
                for linha in linhas_endereco:
                    linhas_quebradas = textwrap.wrap(linha, width=largura_maxima_x // 9)

                    for linha_quebrada in linhas_quebradas:
                        if coordenada_endereco[1] < modelo_verso.size[1]:
                            desenho_verso.text(coordenada_endereco, linha_quebrada, fill=(0, 0, 0), font=fonte_endereco)
                            coordenada_endereco = (coordenada_endereco[0], coordenada_endereco[1] + 40)  # Aumente o valor (ex: +50) conforme necessário
                        else:
                            break

                coordenada_atv2 = (coordenadas_verso["ATV2"][0], coordenadas_verso["ATV2"][1])

                



                desenhar_texto_quebrado(coordenadas_verso["END"], endereco, desenho_verso, fonte_endereco, largura_maxima_x)
                desenhar_texto_quebrado(coordenadas_verso["ATV1"], cnae1 + " - " + atv1, desenho_verso, fonte, largura_maxima_x)
                if atv2:
                    desenhar_texto_quebrado(coordenadas_verso["ATV2"], cnae2 + " - " + atv2, desenho_verso, fonte, largura_maxima_x)
                desenhar_texto_quebrado(coordenadas_verso["LOC"], latitude + "  " + longitude, desenho_verso, fonte, largura_maxima_x)

                
        # Salvar o novo PNG
                modelo.save(cpf + ".pdf")
                modelo_verso.save("verso.pdf")
                # Nome dos arquivos PDF que você gerou
                pdf1_filename = cpf + ".pdf"
                pdf2_filename = "verso.pdf"

                # Abra os arquivos PDF
                pdf1 = PyPDF2.PdfReader(pdf1_filename)
                pdf2 = PyPDF2.PdfReader(pdf2_filename)

                # Crie um objeto PDFFileMerger para mesclar os PDFs
                pdf_merger = PyPDF2.PdfMerger()

                # Adicione os arquivos PDF à mesclagem
                pdf_merger.append(pdf1)
                pdf_merger.append(pdf2)

                output_pdf_filename = os.path.join(r"I:\CARTEIRAS DIGITAIS",unloc_da_pagina, pasta_data_atual , nome  + '.pdf')
                pdf_merger.write(output_pdf_filename)

                # Feche o arquivo PDF de saída
                pdf_merger.close()

                # Excluir os arquivos originais (frente e verso)
                os.remove(cpf + ".pdf")
                os.remove("verso.pdf")

                # Após a geração do PDF com sucesso, você pode remover a linha correspondente ao CPF do DataFrame
                planilha = planilha[planilha['CPF'] != cpf]

                break
        workbook.save('banco_de_dados.xlsx')
        # Salvar a planilha atualizada após a exclusão do CPF
        planilha.to_excel('cpfs.xlsx', index=False)

        break

def memorando():
    from datetime import datetime

    # Mapeamento de municípios para UNLOC
    mapeamento_unloc = {
        'Alvaraes': 'ALV',
        'Açuanopolis':'PAR',
        'Amatura': 'AMT',
        'Anamã': 'ANA',
        'Anori': 'ANO',
        'Apuí': 'APU',
        'Atalaia do Norte': 'ATN',
        'Autazes': 'ATZ',
        'Barcelos': 'BAZ',
        'Balbina':'BNA',
        'Barreirinha': 'BAR',
        'Barreirinha': 'BAE',
        'Benjamin Constant': 'BJC',
        'Beruri': 'BER',
        'Boa Vista do Ramos': 'BVR',
        'Boca do Acre': 'BOA',
        'Borba': 'BBA',
        'Caapiranga': 'CAP',
        'Canutama': 'CAN',
        'Carauari': 'CAF',
        'Careiro': 'CAR',
        'Careiro da Várzea': 'CAZ',
        'Coari': 'CRZ',
        'Coari': 'CIZ',
        'Codajás': 'COD',
        'Eirunepé': 'ERN',
        'Envira': 'ENV',
        'Fonte Boa': 'FBA',
        'Guajará': 'GAJ',
        'Humaitá': 'HIA',
        'Ipixuna': 'IPX',
        'Iranduba': 'IRB',
        'Itamarati': 'ITA',
        'Itacoatiara': 'ITR',
        'Itapiranga': 'ITG',
        'Japurá': 'JPR',
        'Juruá': 'JUR',
        'Jutaí': 'JUT',
        'Lábrea': 'LBR',
        'Manacapuru': 'MPU',
        'Manaquiri': 'MQR',
        'Manaus': 'MAO',
        'Monte Sinai': 'ATZ-MTS',
        'Monte Sinai': 'MTS',
        'Manicoré': 'MNX',
        'Maraã': 'MRA',
        'Maués': 'MBZ',
        'Nhamundá': 'NMD',
        'Nova Olinda do Norte': 'NON',
        'Novo Airão': 'NAR',
        'Novo Aripuanã': 'NAP',
        'Novo Remanso': 'NRO',
        'Parintins': 'PAR',
        'Pauini': 'PUI',
        'Presidente Figueiredo': 'PRF',
        'Rio Preto da Eva': 'RPE',
        'Santa Isabel do Rio Negro': 'SIR',
        'Santo Antônio do Içá': 'SAI',
        'Santo Antonio do Matupi': 'MTP',
        'São Gabriel da Cachoeira': 'SJL',
        'São Paulo de Olivenca': 'SPO',
        'São Sebastião do Uatumã': 'SSU',
        'Sul de Canutama': 'SUL-CAN',
        'Silves': 'SLV',
        'Tabatinga': 'TBT',
        'Tapauá': 'TPA',
        'Tefé': 'TFF',
        'Tonantins': 'TNT',
        'Uarini': 'UAN',
        'Urucará': 'URC',
        'Urucurituba': 'UCB',
        'Vila Extrema': 'VE',
        'Vila Rica de Caviana':'VRC',
        'Vila de Lindoia':'VLD',
        'Vila Realidade':'RLD',
    }

    # Abre o arquivo Excel
    workbook = openpyxl.load_workbook(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")

    # Carregar o modelo do Word
    doc = Document(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\MemorandoSaida\modelo_memosaida.docx")

    # Carregar os dados do Excel em um DataFrame
    df = pd.read_excel(r"I:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")

    # Obter entrada do usuário para o município e a data desejados
    memorando = input("Digite o numero do memorando desejado: ")
    municipio = input("Digite o nome do município desejado: ")
    data = input("Digite a data desejada (no formato dd/mm/yyyy): ")
    memo = input ("Digite os Memorandos utilizados: ")

    # Filtrar os dados do DataFrame com base na entrada do usuário
    selected_data = df[(df['municipio'] == municipio) & (df['data'] == data)]

    # Verificar se existem dados para o município e a data especificados
    if selected_data.empty:
        print(f"Não foram encontrados dados para o município '{municipio}' na data '{data}'.")
    else:
        # Ordenar a coluna 'nomes' em ordem alfabética
        selected_data = selected_data.sort_values(by='nomes')
        # Preencher o modelo do Word com as informações coletadas e adicionar numeração
        for paragraph in doc.paragraphs:
            text = paragraph.text

            if "(num)" in text:
                # Adicionar a data formatada somente se a tag "(data)" estiver presente
                text = text.replace("(num)", memorando)
            
            # Verificar se a tag "(data)" está presente no texto
            if "(data)" in text:
                # Adicionar a data formatada somente se a tag "(data)" estiver presente
                text = text.replace("(data)", data)
            
            # Verificar se a tag "(muni)" está presente no texto
            if "(muni)" in text:
                # Substituir o nome do município e aplicar formatação em negrito
                text = text.replace("(muni)", municipio)
                run = paragraph.add_run(municipio)
                run.bold = True  # Colocar o texto em negrito
                paragraph.clear()

            if "(memos)" in text:
                # Adicionar a data formatada somente se a tag "(data)" estiver presente
                text = text.replace("(memos)", memo)
            
            nomes = selected_data['nomes'].tolist()
            # Adicionar numeração antes de cada nome
            nomes_numerados = [f"{i + 1}. {nome}" for i, nome in enumerate(nomes)]
            text = text.replace("(nomes)", '\n'.join(nomes_numerados))

            
            # Obter o total de nomes (último valor da enumeração) e substituir na variável 'qtda'
            qtda = len(nomes_numerados)
            text = text.replace("(qtda)", str(qtda))
            
            paragraph.clear()  # Limpar o parágrafo original
            paragraph.add_run(text)  # Adicionar o texto modificado

        # Imprimir os nomes com quebras de linha
        print('\n'.join(nomes))

        input("Os nomes estão Corretos? Aperte Enter para continuar!")
            
        # Salvar o documento preenchido na pasta desejada
        output_path = r'I:\MEMORANDOS CPCPR\MEMORANDO 2023\memo-dig'
        doc.save(output_path + '\\' + memorando + '-' + municipio + '-' + 'DIG' + '.docx')

            # Obtendo a UNLOC correspondente ao município
    municipio = mapeamento_unloc.get(municipio, '')

    # Verificar se a UNLOC correspondente foi encontrada
    if municipio:
        # Verifique se a pasta correspondente à UNLOC já existe
        pasta_unloc = os.path.join(r"I:\CARTEIRAS DIGITAIS", municipio)

        # Crie uma pasta com o dia atual (formato: YYYY-MM-DD)
        data_formatada = datetime.strptime(data, "%d/%m/%Y").strftime("%d.%m.%Y")

        # Caminho completo para salvar o documento preenchido na pasta da data atual
        output_path = os.path.join(pasta_unloc, data_formatada, memorando + '-' + municipio + '-' + 'DIG' + '.docx')

        # Salvar o documento preenchido na pasa correspondente
        doc.save(output_path)
    else:
        print(f"A UNLOC correspondente ao município '{municipio}' não foi encontrada.")


    # Fechar o arquivo Excel
    workbook.close()



# Menu principal
while True:
    print("Escolha uma opção:")
    print("1. Inserir CPFs")
    print("2. Emitir Cartões")
    print("3. Emitir Cartão Solo")
    print("4. Fazer Memorando")
    print("5. Trocar Senha")
    print("6. Sair")

    escolha = input("Opção: ")
    if escolha == '1':
        inserir_cpfs()
    elif escolha == '2':
        emitir_cartoes()
    elif escolha == '3':
        emitir_cartoes_slash()
    elif escolha == '4':
        memorando()
    elif escolha == '5':
        nova_senha = input("Digite a nova senha: ")
        salvar_nova_senha(nova_senha)
        print("Senha atualizada com sucesso.")
    elif escolha == '6':
        print("Encerrando o programa.")
        break
    
       
    else:
        print("Opção inválida. Escolha 1, 2, 3, 4, 5 OU 6.")
